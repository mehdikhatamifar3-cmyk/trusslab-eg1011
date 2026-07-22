from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, Optional

import pandas as pd
from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY = "123B5D"
BLUE = "0B4F8A"
LIGHT_BLUE = "DDEBF7"
PALE_BLUE = "EEF5FB"
GREY = "5B6573"
LIGHT_GREY = "F3F5F7"
WHITE = "FFFFFF"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def _add_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "EG1011 Statics and Dynamics | Practical - Forces in a Plane Truss"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor.from_string(GREY)
        _add_page_number(section.footer.paragraphs[0])


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04

    for style_name, size, colour in (
        ("Title", 23, NAVY),
        ("Heading 1", 15, BLUE),
        ("Heading 2", 11.5, NAVY),
        ("Heading 3", 10.5, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Heading 3" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)

    _add_header_footer(document)


def _add_title_block(document: Document, data: Dict[str, Any]) -> None:
    logo_bytes = data.get("jcu_logo")
    if logo_bytes:
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(2.85))

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Practical Report\n")
    subtitle = title.add_run("Forces in a Plane Truss")
    subtitle.font.size = Pt(17)
    subtitle.font.color.rgb = RGBColor.from_string(BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EG1011 Statics and Dynamics")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GREY)

    designer = data.get("designer_name", "")
    if designer:
        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p2.add_run(f"TrussLab practical app designed by {designer}")
        rr.italic = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor.from_string("E8E8E8")

    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    fields = [
        ("Student name", data.get("student_name", "")),
        ("Student ID", data.get("student_id", "")),
        ("Practical/tutorial group", data.get("group", "")),
        ("Practical mode", data.get("mode", "")),
    ]
    for row, (label, value) in zip(table.rows, fields):
        row.cells[0].text = label
        row.cells[1].text = str(value)
        _set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated = data.get("generated_at") or datetime.now().strftime("%d %B %Y, %H:%M")
    run = p.add_run(f"Template generated by TrussLab on {generated}")
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GREY)

    note = document.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    note.cell(0, 0).text = (
        "Student instruction: TrussLab has inserted your identifying details, recorded data, numerical answers, "
        "tables and figures. Complete every shaded writing box in your own words. Keep the headings, margins, "
        "page breaks and box sizes unchanged so all submissions remain consistent in structure and length. "
        "Review the completed eight-page DOCX and submit it through the Practical submission link in LearnJCU."
    )
    _set_cell_shading(note.cell(0, 0), PALE_BLUE)
    _set_cell_margins(note.cell(0, 0), top=130, bottom=130)


def _add_picture(document: Document, image_bytes: Optional[bytes], caption: str, width_inches: float) -> None:
    if not image_bytes:
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.3)
    r.font.color.rgb = RGBColor.from_string(GREY)


def _add_dataframe_table(
    document: Document,
    dataframe: pd.DataFrame,
    formats: Optional[Dict[str, str]] = None,
    font_size: float = 8.2,
) -> None:
    df = dataframe.copy()
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    _set_repeat_table_header(hdr)
    for idx, col in enumerate(df.columns):
        cell = hdr.cells[idx]
        cell.text = str(col)
        _set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size = Pt(font_size)
        _set_cell_margins(cell, top=70, bottom=70)

    for row_idx, (_, series) in enumerate(df.iterrows()):
        row = table.add_row()
        for idx, col in enumerate(df.columns):
            value = series[col]
            if pd.isna(value):
                text = ""
            elif formats and col in formats:
                text = formats[col].format(value)
            else:
                text = str(value)
            cell = row.cells[idx]
            cell.text = text
            _set_cell_margins(cell, top=65, bottom=65)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_size)
        if row_idx % 2 == 1:
            for cell in row.cells:
                _set_cell_shading(cell, "F7F9FB")


def _add_response_box(
    document: Document,
    prompt: str,
    height_cm: float,
    placeholder: str = "Click in this box and type your response. Do not resize the box.",
) -> None:
    table = document.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cell = table.cell(0, 0)
    header_cell.text = prompt
    _set_cell_shading(header_cell, LIGHT_BLUE)
    _set_cell_margins(header_cell, top=80, bottom=80)
    for run in header_cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9.2)
        run.font.color.rgb = RGBColor.from_string(NAVY)

    response_row = table.rows[1]
    response_row.height = Cm(height_cm)
    response_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    response_cell = response_row.cells[0]
    response_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_shading(response_cell, LIGHT_GREY)
    _set_cell_margins(response_cell, top=120, start=130, bottom=100, end=130)
    response_cell.text = placeholder
    run = response_cell.paragraphs[0].runs[0]
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GREY)

    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_instruction(document: Document, text: str) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.8)
    r.font.color.rgb = RGBColor.from_string(GREY)


def build_practical_report(data: Dict[str, Any]) -> bytes:
    """Build a uniform, editable student report template and return DOCX bytes."""
    document = Document()
    _configure_document(document)
    _add_title_block(document, data)

    # Page 2 - student-written introduction and method.
    document.add_page_break()
    document.add_heading("1. Introduction and Aim", level=1)
    _add_instruction(
        document,
        "In approximately 120-160 words, explain why trusses are important, identify the main statics concepts investigated and state a clear aim for this practical.",
    )
    _add_response_box(document, "Student-written introduction and aim", 4.1)

    document.add_heading("2. Apparatus and Method", level=1)
    mode_note = data.get("mode_note", "")
    if mode_note:
        p = document.add_paragraph()
        p.add_run("Practical mode recorded by the app: ").bold = True
        p.add_run(mode_note)
    _add_picture(document, data.get("truss_image"), "Figure 1. Plane truss used in the practical.", 4.1)
    _add_response_box(
        document,
        "Describe the apparatus and summarise the procedure in your own words (approximately 100-140 words)",
        4.0,
    )

    # Page 3 - app-entered responses and experimental data.
    document.add_page_break()
    document.add_heading("3. Preliminary Responses", level=1)
    prelab = data.get("prelab_responses")
    if isinstance(prelab, pd.DataFrame):
        document.add_heading("3.1 Pre-lab responses", level=2)
        _add_dataframe_table(document, prelab, font_size=8.0)
    predictions = data.get("predictions")
    if isinstance(predictions, pd.DataFrame):
        document.add_heading("3.2 Member predictions", level=2)
        _add_dataframe_table(document, predictions, font_size=8.0)

    document.add_heading("4. Experimental or Simulated Data", level=1)
    document.add_paragraph(
        "Sign convention: positive values indicate tension and negative values indicate compression. Dial displacement is reported in millimetres."
    )
    lab_data = data.get("lab_data")
    if isinstance(lab_data, pd.DataFrame):
        _add_dataframe_table(
            document,
            lab_data,
            formats={"10 kg (mm)": "{:+.3f}", "20 kg (mm)": "{:+.3f}", "30 kg (mm)": "{:+.3f}"},
            font_size=8.2,
        )
    p = document.add_paragraph()
    p.add_run("Return-to-zero observation: ").bold = True
    p.add_run(str(data.get("zero_return", "")))
    _add_response_box(
        document,
        "Brief data-quality observation: identify any unusual reading, zero drift or limitation noticed during data collection",
        2.1,
    )

    # Page 4 - calculations.
    document.add_page_break()
    document.add_heading("5. Theoretical Calculations", level=1)
    document.add_paragraph(
        "The following numerical values were entered in TrussLab for the 30 kg load case. Positive member force indicates tension; negative member force indicates compression."
    )
    calc_df = pd.DataFrame(
        [
            ["Applied load, W", data.get("calc_w"), "N"],
            ["Vertical reaction, Ay", data.get("calc_ay"), "N"],
            ["Member force, F_AE", data.get("calc_ae"), "N"],
            ["Member force, F_AB", data.get("calc_ab"), "N"],
            ["Member force, F_EB", data.get("calc_eb"), "N"],
        ],
        columns=["Quantity", "Student value", "Unit"],
    )
    _add_dataframe_table(document, calc_df, formats={"Student value": "{:+.2f}"}, font_size=8.5)
    _add_instruction(
        document,
        "Show sufficient working for W = mg, the support reactions and joint A. Include free-body reasoning, equilibrium equations, substitutions, units and tension/compression interpretation.",
    )
    _add_response_box(document, "Detailed theoretical calculation working", 14.1)

    # Page 5 - results and analysis.
    document.add_page_break()
    document.add_heading("6. Results and Analysis", level=1)
    comparison = data.get("comparison")
    if isinstance(comparison, pd.DataFrame):
        _add_dataframe_table(
            document,
            comparison,
            formats={
                "Dial (mm)": "{:+.3f}",
                "Calibration (N/mm)": "{:.1f}",
                "Experimental force (N)": "{:+.2f}",
                "Student theoretical force (N)": "{:+.2f}",
            },
            font_size=7.3,
        )
    experimental_value = data.get("experimental_calc")
    experimental_text = "No value entered" if experimental_value is None else f"{float(experimental_value):+.2f} N"
    p = document.add_paragraph()
    p.add_run("Student-entered worked experimental force for member AE at 30 kg: ").bold = True
    p.add_run(experimental_text)
    _add_picture(document, data.get("dial_graph"), "Figure 2. Dial displacement as the applied mass increased.", 5.25)
    _add_response_box(
        document,
        "Describe the important trends and compare the experimental/simulated forces with the theoretical results (approximately 120-160 words)",
        4.8,
    )

    # Page 6 - guided discussion.
    document.add_page_break()
    document.add_heading("7. Discussion", level=1)
    _add_instruction(document, "Answer each question using evidence from your tables, calculations and graph.")
    prompts = [
        "7.1 Which instrumented members were in tension, compression and approximately zero force? Explain briefly.",
        "7.2 Did dial displacement change approximately in proportion to applied load? Support your answer with one numerical example.",
        "7.3 Why might a theoretical zero-force member show a small measured or simulated force? Identify relevant practical assumptions or errors.",
        "7.4 Which measured member force was closest to your theoretical value at 30 kg? State the evidence used.",
    ]
    for prompt in prompts:
        _add_response_box(document, prompt, 3.25)

    # Page 7 - Part B.
    document.add_page_break()
    document.add_heading("8. Part B - Safe Load Engineering Challenge", level=1)
    _add_picture(document, data.get("safe_load_image"), "Figure 3. Member-force utilisation at the selected applied mass.", 5.0)
    part_b_df = pd.DataFrame(
        [
            ["Maximum safe mass", data.get("safe_mass_answer"), "kg"],
            ["Controlling member", data.get("critical_member_answer"), ""],
            ["Controlling force type", data.get("critical_type_answer"), ""],
        ],
        columns=["Design response", "Student answer", "Unit"],
    )
    _add_dataframe_table(document, part_b_df, font_size=8.5)
    _add_response_box(
        document,
        "Show how the limiting mass was determined from the tensile and compressive allowable values. Explain why the 30 kg laboratory limit is reasonable.",
        6.0,
    )

    # Page 8 - conclusion, references and declaration.
    document.add_page_break()
    document.add_heading("9. Conclusion", level=1)
    _add_instruction(
        document,
        "In approximately 80-120 words, state the main findings, whether the practical aim was achieved and the most important limitation or improvement.",
    )
    _add_response_box(document, "Student-written conclusion", 5.0)

    document.add_heading("10. References", level=1)
    _add_instruction(document, "List the practical handout and any other sources used. Apply one consistent referencing style.")
    _add_response_box(document, "References", 2.8)

    document.add_heading("11. Student Review Declaration", level=1)
    declaration = document.add_table(rows=1, cols=1)
    declaration.style = "Table Grid"
    declaration.cell(0, 0).text = (
        "I confirm that I reviewed the app-entered data and completed the written sections of this report in my own words. "
        "I did not change the prescribed headings, margins, page breaks or response-box sizes."
    )
    _set_cell_shading(declaration.cell(0, 0), PALE_BLUE)
    _set_cell_margins(declaration.cell(0, 0), top=130, bottom=130)

    p = document.add_paragraph()
    p.add_run("Student name: ").bold = True
    p.add_run(str(data.get("student_name", "")))
    p.add_run("\nDate reviewed: ____________________")
    p.add_run("\nStudent signature: ______________________________")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
